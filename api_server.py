import os
from io import BytesIO
from typing import Optional

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from docx import Document

from main import run_model
from main import warmup
warmup()
app = FastAPI(title="ML Chat API")

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def read_docx_bytes(file_bytes: bytes) -> str:
    """
    Читает текст из .docx из bytes.
    """
    doc = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                parts.append(" | ".join(row_parts))

    return "\n".join(parts).strip()


@app.get("/api/health")
async def health():
    return {"ok": True, "status": "running"}


@app.post("/api/chat")
async def chat(
    user_text: str = Form(default=""),
    docx_file: Optional[UploadFile] = File(default=None),
):
    docx_text = None
    filename = None

    if docx_file is not None:
        filename = docx_file.filename or "unknown"

        if not filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Разрешены только .docx файлы")

        file_bytes = await docx_file.read()
        docx_text = read_docx_bytes(file_bytes)

    if not user_text.strip() and not docx_text:
        raise HTTPException(status_code=400, detail="Нужно передать текст или .docx файл")

    answer = run_model(user_text=user_text.strip(), docx_text=docx_text)

    return {
        "ok": True,
        "answer": answer,
        "meta": {
            "filename": filename,
            "docx_preview": docx_text[:500] if docx_text else "",
            "user_text_length": len(user_text.strip()),
        },
    }